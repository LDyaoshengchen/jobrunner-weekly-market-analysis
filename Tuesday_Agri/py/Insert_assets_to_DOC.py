# Insert_assets_to_DOC.py (Tuesday - simplified)
# Insert only:
#  - Wheat price chart
#  - Wheat CFTC net position chart

from pathlib import Path
from docx import Document
from docx.shared import Inches
from datetime import datetime

# =====================================
# Paths
# =====================================
CURRENT_FILE = Path(__file__).resolve()
ROOT_DIR = CURRENT_FILE.parents[2]

ASSETS_ROOT = ROOT_DIR / "assets"
ASSETS_TUE = ROOT_DIR / "Tuesday_Agri" / "assets"

TEMPLATE_PATH = ASSETS_TUE / "template_Agriculture.docx"

OUTPUT_DIR = ROOT_DIR / "Tuesday_Agri" / "output"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

today_str = datetime.now().strftime("%Y%m%d")
OUTPUT_PATH = OUTPUT_DIR / f"Soft_Commodities_Report_Wheat.docx"

# =====================================
# Inputs (fixed filenames)
# =====================================
PRICE_CHART = ASSETS_TUE / "price_chart_Wheat.png"
CFTC_NETPOS = ASSETS_ROOT / "cftc_net_position_Wheat.png"

IMAGE_MAP = {
    "<<FIG_WHEAT_CHART>>": PRICE_CHART,
    "<<FIG_WHEAT_NETPOSITION>>": CFTC_NETPOS,
}

# =====================================
# Insert images into template
# =====================================
def insert_images(doc: Document, image_map: dict[str, Path], width_inch: float = 5.5):
    for p in doc.paragraphs:
        for tag, img_path in image_map.items():
            if tag in p.text:
                p.text = p.text.replace(tag, "")
                if img_path.exists():
                    run = p.add_run()
                    run.add_picture(str(img_path), width=Inches(width_inch))
                    print(f"[OK] Inserted: {img_path.name}")
                else:
                    print(f"[WARN] Missing file: {img_path}")

def main():
    if not TEMPLATE_PATH.exists():
        print(f"[ERROR] Template not found: {TEMPLATE_PATH}")
        return

    doc = Document(str(TEMPLATE_PATH))
    insert_images(doc, IMAGE_MAP)
    doc.save(str(OUTPUT_PATH))

    print(f"[DONE] Saved: {OUTPUT_PATH}")

if __name__ == "__main__":
    main()
