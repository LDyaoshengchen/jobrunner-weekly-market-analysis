# Insert_assets_to_DOC.py (Thursday - simplified)
# Insert only:
#  - Crude Oil price chart
#  - Crude Oil CFTC net position chart

from pathlib import Path
from docx import Document
from docx.shared import Inches

# =====================================
# Paths
# =====================================
CURRENT_FILE = Path(__file__).resolve()
ROOT_DIR = CURRENT_FILE.parents[2]

ASSETS_ROOT = ROOT_DIR / "assets"
ASSETS_THU = ROOT_DIR / "Thursday_Energy" / "assets"

TEMPLATE_PATH = ASSETS_THU / "template_Energy.docx"

OUTPUT_DIR = ROOT_DIR / "Thursday_Energy" / "output"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

OUTPUT_PATH = OUTPUT_DIR / "Energy_Report_Crude_Oil.docx"

# =====================================
# Inputs (fixed filenames for clarity)
# =====================================
PRICE_CHART = ASSETS_THU / "price_chart_Crude_Oil.png"
CFTC_NETPOS  = ASSETS_ROOT / "cftc_net_position_Crude_Oil.png"

IMAGE_MAP = {
    "<<FIG_Crude_Oil_CHART>>": PRICE_CHART,
    "<<FIG_Crude_Oil_NETPOSITION>>": CFTC_NETPOS,
}

# =====================================
# Insert images into template
# =====================================
def insert_images(doc: Document, image_map: dict[str, Path], width_inch: float = 6.0):
    for p in doc.paragraphs:
        for tag, img_path in image_map.items():
            if tag in p.text:
                # Remove tag text
                p.text = p.text.replace(tag, "")
                # Insert image if exists
                if img_path.exists():
                    run = p.add_run()
                    run.add_picture(str(img_path), width=Inches(width_inch))
                    print(f"[OK] Inserted: {img_path.name}")
                else:
                    print(f"[WARN] Missing file: {img_path}")

def main():
    if not TEMPLATE_PATH.exists():
        print(f"[ERROR] Template not found: {TEMPLATE_PATH}")
        return

    doc = Document(str(TEMPLATE_PATH))
    insert_images(doc, IMAGE_MAP, width_inch=6.0)
    doc.save(str(OUTPUT_PATH))
    print(f"[DONE] Saved: {OUTPUT_PATH}")

if __name__ == "__main__":
    main()
